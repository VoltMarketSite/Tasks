import os, fitz, json
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from openpyxl import load_workbook


SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt", ".xlsx", ".jpg", ".jpeg", ".png"]

path = os.path.dirname(os.path.abspath(__file__))
OUTPUT_IMAGE_DIR = f"{path}\\temp\\extracted_images"
os.makedirs(OUTPUT_IMAGE_DIR, exist_ok=True)

def extract_text_from_docx(path):
    try:
        doc = Document(path)
        full_text = []

        # 1. Обычные параграфы
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text) 

        # 2. Таблицы
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)

    except Exception as e:
        print(f"❌ Ошибка при чтении DOCX: {e}")
        return ""

def extract_text_from_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text_from_pdf(path):
    try:
        doc = fitz.open(path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        print(f"Ошибка чтения PDF: {e}")
        return ""

def extract_images_from_pdf(path):
    try:
        images = convert_from_path(path, dpi=300, poppler_path="C:\\Program Files\\poppler-24.08.0\\Library\\bin")
        saved_paths = []
        if not os.path.exists(OUTPUT_IMAGE_DIR):
            os.mkdir(OUTPUT_IMAGE_DIR)
        for i, img in enumerate(images):
            img_path = os.path.join(OUTPUT_IMAGE_DIR, f"{os.path.basename(path)}_page_{i+1}.png")
            img.save(img_path, "PNG")
            saved_paths.append(img_path)
        return saved_paths
    except Exception as e:
        print(f"Ошибка при извлечении изображений из PDF: {e}")
        return []

def extract_text_from_xlsx(path):
    wb = load_workbook(path)
    text = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            row_text = []
            for cell in row:
                if cell.value is not None:
                    row_text.append(str(cell.value))
            if row_text:
                text.append("\t".join(row_text))
    return "\n".join(text)

def extract_file_content(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return {"text": extract_text_from_docx(path), "images": []}
    elif ext == ".txt":
        return {"text": extract_text_from_txt(path), "images": []}
    elif ext == ".xlsx":
        return {"text": extract_text_from_xlsx(path), "images": []}
    elif ext == ".pdf":
        text = extract_text_from_pdf(path)
        images = extract_images_from_pdf(path)
        return {"text": text, "images": images}
    elif ext in [".jpg", ".jpeg", ".png"]:
        return {"text": "", "images": [path]}
    else:
        return {"text": "", "images": []}

def extract_json(text):
    start = text.find('{')
    if start == -1:
        return None

    brace_count = 0
    for i in range(start, len(text)):
        if text[i] == '{':
            brace_count += 1
        elif text[i] == '}':
            brace_count -= 1

        if brace_count == 0:
            try:
                json_str = text[start:i + 1]
                return json.loads(json_str)
            except json.JSONDecodeError:
                return None
    return None

def extracting():
    all_data = {"texts": [], "images": []}
    q = os.listdir(f"{path}\\temp")
    for file in q:
        file_abs = f"{path}\\temp\\{file}"
        if os.path.isdir(file_abs): continue
        data = extract_file_content(file_abs)
        all_data["texts"].append(data["text"])
        all_data["images"].extend(data["images"])
    
    return all_data